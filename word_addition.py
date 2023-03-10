import datetime

from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement, ns
from docx.shared import Cm
from docx.shared import Inches
from docx.shared import Pt


def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def add_page_number(paragraph):
    """Метод, добавляющий номера страниц в word."""
    # выравниваем параграф по центру
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # запускаем динамическое обновление параграфа
    page_num_run = paragraph.add_run()
    # обозначаем начало позиции вывода
    fld_char1 = create_element('w:fldChar')
    create_attribute(fld_char1, 'w:fldCharType', 'begin')
    # задаем вывод текущего значения страницы PAGE (всего страниц NUM PAGES)
    instr_text = create_element('w:instrText')
    create_attribute(instr_text, 'xml:space', 'preserve')
    instr_text.text = "PAGE"
    # обозначаем конец позиции вывода
    fld_char2 = create_element('w:fldChar')
    create_attribute(fld_char2, 'w:fldCharType', 'end')
    # добавляем все в наш параграф (который формируется динамически)
    page_num_run._r.append(fld_char1)
    page_num_run._r.append(instr_text)
    page_num_run._r.append(fld_char2)


def page_breaker(doc):
    run = doc.add_paragraph().add_run()
    run.add_break(WD_BREAK.PAGE)


def add_new_styles(document, style_name: str, font_size: int, bold: bool = False, italic: bool = False):
    """Метод, добавляющий стили текста."""
    styles = document.styles
    styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
    style = document.styles[style_name]
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(font_size)
    font.bold = bold
    font.italic = italic


def section_choice(document):
    """Метод, добавляющий отступы в документе word."""
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)


def make_table_bold(table, cols, rows):
    """Метод, изменяющий вес шрифтов в таблицах и выравнивающий таблицу по центру."""
    for row in range(1):
        for col in range(cols):
            # получаем ячейку таблицы
            cell = table.cell(row, col)
            # записываем в ячейку данные
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True

    for row in range(1, rows):
        for col in range(1):
            # получаем ячейку таблицы
            cell = table.cell(row, col)
            # записываем в ячейку данные
            run = cell.paragraphs[0].runs[0]
            run.font.bold = True
    table.alignment = WD_TABLE_ALIGNMENT.CENTER


def change_cell_size(table, column_num, size_arr):
    """Метод, меняющий размер клеток в таблице."""
    for i in range(column_num):
        for cell in table.columns[i].cells:
            cell.width = Inches(size_arr[i])


def adding_graphic(document, title, picture_path, width: float, height: float):
    """Метод, добавляющий в word график."""
    document.add_picture(picture_path, width=Inches(width), height=Inches(height))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    picture_title = document.add_paragraph(title, style='PItalic')
    picture_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


def time_breaks_counter(break_frame):
    fail_str_begin = []
    fail_str_end = []
    lost_minutes = []
    for i in range(len(break_frame.index)):
        fail_str_begin.append(f"{break_frame['StartDate'][i]}  {break_frame['StartTime'][i]}")
        fail_str_end.append(f"{break_frame['EndDate'][i]}  {break_frame['EndTime'][i]}")
        lost_minutes.append(round((datetime.datetime.combine(break_frame['EndDate'][i],
                                                             break_frame['EndTime'][i]) -
                                   datetime.datetime.combine(break_frame['StartDate'][i],
                                                             break_frame['StartTime'][i])).total_seconds() / 60, 2))
    return fail_str_begin, fail_str_end, lost_minutes
