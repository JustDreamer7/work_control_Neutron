import time
import datetime
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from collections import defaultdict
import word_addition
from graph_drawer import GraphsDrawing
from proccessing_data_neutron import ProccessingNeutron


def make_report_neutron(start_date, end_date, report_path, picture_path, neutron_data, pressure_data,
                        proccessing_pressure, accessory_data):
    t1 = time.time()

    days_amount = len(pd.date_range(start_date, end_date))

    processing_inst = ProccessingNeutron(df_data=neutron_data, pressure_data=pressure_data,
                                         default_pressure=proccessing_pressure)

    graphs = GraphsDrawing(start_date=start_date, end_date=end_date,
                           path_to_pic=f'{picture_path}')

    graphs.change_design()

    worktime_frame, break_frame, parameters_dict = processing_inst.period_processing_for_report(
        start_date=start_date, end_date=end_date)

    n_pressure_graph = graphs.pressure_graph(corr_pressure_data=parameters_dict['corr_pressure_n'],
                                             neutron_data=neutron_data,
                                             fit_line=parameters_dict['fit_line'],
                                             type_of_impulse='Nn')
    noise_pressure_graph = graphs.pressure_graph(corr_pressure_data=parameters_dict['corr_pressure_noise'],
                                                 neutron_data=neutron_data,
                                                 fit_line=parameters_dict['fit_line_noise'],
                                                 type_of_impulse='N_noise')
    neutron_graph = graphs.neutron_graph(neutron_data=neutron_data,
                                         corr_for_neutron_data=parameters_dict['correction_for_n'],
                                         pressure_data=pressure_data,
                                         type_of_impulse='Nn')
    noise__graph = graphs.neutron_graph(neutron_data=neutron_data,
                                        corr_for_neutron_data=parameters_dict['correction_for_noise'],
                                        pressure_data=pressure_data,
                                        type_of_impulse='N_noise')
    mask_event_dict = defaultdict(list)
    for det in range(1, 5):
        n_check_df = neutron_data[neutron_data[f'Nn{det}'] > np.median(neutron_data[f'Nn{det}']) * 1.5]
        noise_check_df = neutron_data[neutron_data[f'N_noise{det}'] > np.median(neutron_data[f'N_noise{det}']) * 1.5]
        for row in list(n_check_df.index):
            mask_event_dict['det'].append(det)
            mask_event_dict['datetime'].append(n_check_df['datetime'][row])
            mask_event_dict['type'].append('Нейтроны')
            mask_event_dict['value'].append(n_check_df[f'Nn{det}'][row])
            mask_event_dict['n/med'].append(round(n_check_df[f'Nn{det}'][row] / np.median(neutron_data[f'Nn{det}']), 2))

        for row in list(noise_check_df.index):
            mask_event_dict['det'].append(det)
            mask_event_dict['datetime'].append(noise_check_df['datetime'][row])
            mask_event_dict['type'].append('Шумы')
            mask_event_dict['value'].append(noise_check_df[f'N_noise{det}'][row])
            mask_event_dict['n/med'].append(
                round(noise_check_df[f'N_noise{det}'][row] / np.median(neutron_data[f'N_noise{det}']), 2))

    r_distribution_data = accessory_data.r_file_reader()

    front_time_data = accessory_data.front_time_data_reader()

    n_amp_data = accessory_data.n_amp_data_reader()

    print(f'{worktime_frame}')
    print(f'{break_frame}')
    print(f'{parameters_dict["N_0"]}')

    # Далее формирование и запись в word-файл

    doc = Document()
    word_addition.section_choice(doc)
    word_addition.add_new_styles(doc, style_name='PItalic', font_size=11, bold=True, italic=True)
    word_addition.add_new_styles(doc, style_name='Head-style', font_size=14, bold=True)
    word_addition.add_new_styles(doc, style_name='Head-graphic', font_size=13, bold=True, italic=True)

    head = doc.add_paragraph(
        f'Справка о работе установки «Нейтрон» в период с {start_date} по {end_date}', style='Head-style')
    head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    table_title = doc.add_paragraph('Таблица 1: Календарная эффективность.', style='PItalic')
    table_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    worktime_table = doc.add_table(5, 4, doc.styles['Table Grid'])
    worktime_table.cell(0, 0).text = '№ детектора'
    worktime_table.cell(0, 1).text = 'Экспозиции, ч.'
    worktime_table.cell(0, 2).text = 'Календарное время, ч.'
    worktime_table.cell(0, 3).text = 'Экспозиция, %'
    for idx in range(1, 5):
        worktime_table.cell(idx, 0).text = f'{idx}'
        worktime_table.cell(idx, 1).text = str(round(worktime_frame[f'Worktime_{idx}'].sum(), 2))
        worktime_table.cell(idx, 2).text = str(24 * days_amount)
        worktime_table.cell(idx, 3).text = str(
            round(worktime_frame[f'Worktime_{idx}'].sum() / (24 * days_amount) * 100, 3)) + '%'

    word_addition.make_table_bold(worktime_table, cols=4, rows=5)
    doc.add_paragraph()

    fail_str_begin, fail_str_end, lost_minutes = word_addition.time_breaks_counter(break_frame=break_frame)

    break_table_title = doc.add_paragraph('Таблица 2: Сводная таблица остановок установки Нейтрон.', style='PItalic')
    break_table_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    break_table = doc.add_table(len(fail_str_begin) + 2, 5, doc.styles['Table Grid'])
    break_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    break_table.cell(0, 0).text = '№'
    break_table.cell(0, 0).merge(break_table.cell(1, 0))
    break_table.cell(0, 1).text = 'Время простоя'
    break_table.cell(1, 1).text = 'c'
    break_table.cell(1, 2).text = 'по'
    break_table.cell(0, 1).merge(break_table.cell(0, 2))
    break_table.cell(0, 3).text = 'Кол-во потерянных минут (период)'
    break_table.cell(0, 3).merge(break_table.cell(1, 3))
    break_table.cell(0, 4).text = 'Примечание'
    break_table.cell(0, 4).merge(break_table.cell(1, 4))

    for i in range(2, len(fail_str_begin) + 2):
        break_table.cell(i, 0).text = str(break_frame['detector'][i - 2])
        break_table.cell(i, 1).text = fail_str_begin[i - 2]
        break_table.cell(i, 2).text = fail_str_end[i - 2]
        break_table.cell(i, 3).text = str(lost_minutes[i - 2])
        break_table.cell(i, 4).text = ' '

    word_addition.make_table_bold(break_table, cols=5, rows=len(fail_str_begin) + 2)
    doc.add_paragraph()

    stat_table_title = doc.add_paragraph(r'Таблица 3: Средние скорости счета [(300с)⁻¹]', style='PItalic')
    stat_table_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    stat_table = doc.add_table(3, 5, doc.styles['Table Grid'])
    stat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    stat_table.cell(0, 0).text = '№ детектора'
    for det in range(1, 5):
        stat_table.cell(0, det).text = f'{det}'
        stat_table.cell(1, det).text = str(round(neutron_data[neutron_data[f'Nn{det}'] != 0][f'Nn{det}'].mean(), 2))
        stat_table.cell(2, det).text = str(
            round(neutron_data[neutron_data[f'N_noise{det}'] != 0][f'N_noise{det}'].mean(), 2))
    stat_table.cell(1, 0).text = 'Скорость счета нейтронных импульсов'
    stat_table.cell(2, 0).text = 'Скорость счета шумовых импульсов'

    word_addition.make_table_bold(stat_table, cols=5, rows=3)
    word_addition.change_cell_size(stat_table, column_num=5, size_arr=[3, 1.075, 1.075, 1.075, 1.075])
    doc.add_paragraph()

    bar_table_title = doc.add_paragraph('Таблица 4: Барометрические коэффициенты β ', style='PItalic')
    bar_table_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    bar_table = doc.add_table(3, 5, doc.styles['Table Grid'])
    bar_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    bar_table.cell(0, 0).text = '№ детектора'
    bar_table.cell(1, 0).text = 'β нейтронных импульсов, %/мбар'
    bar_table.cell(2, 0).text = 'β шумовых импульсов, %/мбар'
    for det in range(1, 5):
        bar_table.cell(0, det).text = f'{det}'
        bar_table.cell(1, det).text = str(round(parameters_dict['N_bar'][det - 1], 2))
        bar_table.cell(2, det).text = str(round(parameters_dict['N_bar_noise'][det - 1], 2))

    word_addition.make_table_bold(bar_table, cols=5, rows=3)
    word_addition.change_cell_size(bar_table, column_num=5, size_arr=[3, 1.075, 1.075, 1.075, 1.075])
    doc.add_paragraph()

    b_table_title = doc.add_paragraph('Таблица 5: Коэффициент B', style='PItalic')
    b_table_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    b_table = doc.add_table(3, 5, doc.styles['Table Grid'])
    b_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    b_table.cell(0, 0).text = '№ детектора'
    b_table.cell(1, 0).text = 'B нейтронных импульсов'
    b_table.cell(2, 0).text = 'B шумовых импульсов'
    for det in range(1, 5):
        b_table.cell(0, det).text = f'{det}'
        b_table.cell(1, det).text = str(round(parameters_dict['B_factor'][det - 1], 2))
        b_table.cell(2, det).text = str(round(parameters_dict['B_factor_noise'][det - 1], 2))

    word_addition.make_table_bold(b_table, cols=5, rows=3)
    word_addition.change_cell_size(b_table, column_num=5, size_arr=[3, 1.075, 1.075, 1.075, 1.075])
    doc.add_paragraph()

    mask_event_table_title = doc.add_paragraph('Таблица 6: События для внесения в маску', style='PItalic')
    mask_event_table_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    mask_event_table = doc.add_table(len(mask_event_dict['det']) + 1, 5, doc.styles['Table Grid'])
    mask_event_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    mask_event_table.cell(0, 0).text = '№ детектора'
    mask_event_table.cell(0, 1).text = 'Время'
    mask_event_table.cell(0, 2).text = 'Тип'
    mask_event_table.cell(0, 3).text = 'Значение'
    mask_event_table.cell(0, 4).text = 'N/N\u2098\u2091'

    for event_idx in range(len(mask_event_dict['det'])):
        mask_event_table.cell(event_idx + 1, 0).text = f'{mask_event_dict["det"][event_idx]}'
        mask_event_table.cell(event_idx + 1, 1).text = f'{mask_event_dict["datetime"][event_idx]}'
        mask_event_table.cell(event_idx + 1, 2).text = f'{mask_event_dict["type"][event_idx]}'
        mask_event_table.cell(event_idx + 1, 3).text = f'{mask_event_dict["value"][event_idx]}'
        mask_event_table.cell(event_idx + 1, 4).text = f'{mask_event_dict["n/med"][event_idx]}'

    word_addition.make_table_bold(mask_event_table, cols=5, rows=len(mask_event_dict['det']) + 1)
    word_addition.change_cell_size(b_table, column_num=5, size_arr=[1.075, 3, 1.075, 1.075, 1.075])
    doc.add_paragraph()

    word_addition.page_breaker(doc=doc)

    graphic_header = doc.add_paragraph('Зависимости скорости счета нейтронных импульсов и давления от времени.',
                                       style='Head-graphic')
    graphic_header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    word_addition.adding_graphic(doc,
                                 title='Рис. 1 - Зависимости скорости счета ' +
                                       'нейтронных импульсов и давления от времени.',
                                 width=7.5,
                                 height=7.9,
                                 picture_path=str(neutron_graph))

    word_addition.page_breaker(doc=doc)

    graphic_header = doc.add_paragraph('Зависимости скорости счета шумовых импульсов и давления от времени.',
                                       style='Head-graphic')
    graphic_header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    word_addition.adding_graphic(doc,
                                 title='Рис. 2 - Зависимости скорости счета ' +
                                       'шумовых импульсов и давления от времени.',
                                 width=7.5,
                                 height=7.9,
                                 picture_path=str(noise__graph))

    word_addition.page_breaker(doc=doc)

    graphic_header = doc.add_paragraph('Зависимости скорости счета нейтронных импульсов от давления.',
                                       style='Head-graphic')
    graphic_header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    word_addition.adding_graphic(doc,
                                 title='Рис. 3 - Зависимости скорости счета нейтронных импульсов от давления.',
                                 width=7.5,
                                 height=7.9,
                                 picture_path=str(n_pressure_graph))

    word_addition.page_breaker(doc=doc)

    graphic_header = doc.add_paragraph('Зависимости скорости счета шумовых импульсов от давления.',
                                       style='Head-graphic')
    graphic_header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    word_addition.adding_graphic(doc,
                                 title='Рис. 4 - Зависимости скорости счета шумовых импульсов от давления.',
                                 width=7.5,
                                 height=7.9,
                                 picture_path=str(noise_pressure_graph))

    word_addition.page_breaker(doc=doc)

    try:
        n_amp_dist_graph = graphs.n_amp_dist(n_amp_frame=n_amp_data,
                                             single_date=accessory_data.single_date)
        graphic_header = doc.add_paragraph('Временные распределения сигналов.',
                                           style='Head-graphic')
        graphic_header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        word_addition.adding_graphic(doc,
                                     title='Рис. 5 - Временные распределения сигналов.',
                                     width=7.5,
                                     height=3.9,
                                     picture_path=str(n_amp_dist_graph))
    except AttributeError:
        print(f"sp-файлов в период за {start_date}/{end_date} не существует")

    try:
        front_time_dist_graph = graphs.front_time_dist(front_time_frame=front_time_data,
                                                       single_date=accessory_data.single_date)
        graphic_header = doc.add_paragraph('Амплитудные распределения сигналов.',
                                           style='Head-graphic')
        graphic_header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        word_addition.adding_graphic(doc,
                                     title='Рис. 6 - Амплитудные распределения сигналов.',
                                     width=7.5,
                                     height=3.9,
                                     picture_path=str(front_time_dist_graph))
    except AttributeError:
        print(f"Tf-файлов в период за {start_date}/{end_date} не существует")

    try:
        r_dist_graph = graphs.r_distribution(r_dist_frame=r_distribution_data,
                                             single_date=accessory_data.single_date)
        graphic_header = doc.add_paragraph('Распределения сигналов по параметру R (Af/Amax).',
                                           style='Head-graphic')
        graphic_header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        word_addition.adding_graphic(doc,
                                     title='Рис. 7 - Распределения сигналов по параметру R (Af/Amax).',
                                     width=7.5,
                                     height=3.9,
                                     picture_path=str(r_dist_graph))
    except AttributeError:
        print(f"R-файлов в период за {start_date}/{end_date} не существует")

    word_addition.add_page_number(doc.sections[0].footer.paragraphs[0])
    doc.save(
        f'{report_path}\\{start_date.day:02}.{start_date.month:02}.{start_date.year}'
        f'-{end_date.day:02}.{end_date.month:02}.{end_date.year}.docx')

    print(time.time() - t1)
